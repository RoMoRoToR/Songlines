"""Refill the four ITMO RID submission documents for the Q/R/M/C package.

This replaces the earlier CADENZA submission. Q/R/M/C is selected as the
РИД unit because it contains substantially stronger novelty markers:

  • Formal definitions with an identifiability argument (Def. 1-2)
  • Empirical Claim 1 verified at p<10⁻⁴ on 35,640 multi-agent runs
  • Substrate-independent Q/R/M/C event logger (MiniGrid + BabyAI + custom)
  • A minimal Collective Semantic Memory that strictly dominates five
    fixed-cadence peer architectures on 3,240 runs
  • Order-theoretic foundation via a Galois connection between query
    intent and retrieval extension (Formal Concept Analysis)

Submitted body of code: experiments/big_experiment/ (2 338 LOC) +
scripts/validate_qrmc_factorization.py (149 LOC) = 2 487 LOC, ~95 KB.

This script does NOT modify document styles, fonts, margins, or chrome.
It only writes content into placeholders. Personal data of the author
(СНИЛС, ИНН, ДР, адрес, телефон, паспорт) remain `[заполнить]`.
"""

from __future__ import annotations

import os
import subprocess

from docx import Document
from docx.shared import Pt


ITMO = "/Users/taniyashuba/PycharmProjects/Songlines/docs/ITMO"
BIG_EXP = "/Users/taniyashuba/PycharmProjects/Songlines/experiments/big_experiment"
SCRIPTS = "/Users/taniyashuba/PycharmProjects/Songlines/scripts"

# ─────────────────────────────────────────────────────────────────────
# Q/R/M/C description (used in multiple documents)
# ─────────────────────────────────────────────────────────────────────

QRMC_NAME = (
    "Q/R/M/C-Диаг — Программный комплекс стадийной декомпозиции и "
    "диагностики систем агентной памяти "
    "(Stage-Decomposition Diagnostic Framework for Memory-Based "
    "Navigation Agents)"
)

QRMC_FULL_DESCRIPTION = (
    "Программный комплекс Q/R/M/C-Диаг реализует оригинальный метод "
    "четырёхстадийной декомпозиции работы агентной памяти при "
    "семантической навигации: Q (формирование запроса), R "
    "(удовлетворение поиска), M (материализация цели), C (завершение "
    "после поиска). Метод включает: (а) формальное определение четырёх "
    "стадий через cтадийно-марковское условие и аргумент об "
    "оценимости конечных частотных оценок из лога траекторий "
    "(Definition 1-2 в сопроводительной статье); (б) расширение на "
    "мультиагентный режим с двумя каналами связности — памяти и "
    "занятости целей (Definition 2); (в) эмпирически подтверждённое "
    "на 35 640 прогонах утверждение о сдвиге узкого места между "
    "стадиями M и C при изменении каденции K (Empirical Claim 1, оба "
    "знака наклонов при p<10⁻⁴, повторено на 100-прогонном "
    "MiniGrid-портабилити-тесте при p<10⁻⁶); (г) субстрат-независимый "
    "логгер Q/R/M/C-событий, единообразно работающий на трёх "
    "окружениях (MiniGrid, BabyAI, кастомная сетка 12×10); "
    "(д) встроенный прототип Collective Semantic Memory — "
    "минимальный экземпляр распределённой коллективной семантической "
    "памяти с явными правилами merge/trust/staleness, статистически "
    "строго доминирующий пять фиксированных peer-broadcast "
    "архитектур на 3 240 прогонах (lower 95% bootstrap CI = 0.610 > "
    "upper CI каждой baseline). Теоретическая новизна: операционные "
    "определения Q/R/M/C получают чтение через антитонное "
    "соответствие Галуа между намерением запроса и экстенсией поиска "
    "(концептуальная решётка по Гантеру-Вилле), что превращает "
    "диагноз «retrieval пуст» в строгое утверждение «семантика "
    "запроса лежит вне реализованной концептуальной решётки». Эта "
    "комбинация формальной структуры, эмпирической верификации на "
    "большой выборке и инструментального применения отсутствует в "
    "существующих архитектурах диагностики мультиагентных систем "
    "памяти и, как предполагается, может составить предмет для "
    "самостоятельной охраны."
)

QRMC_ABSTRACT_900 = (
    "Q/R/M/C-Диаг — Программный комплекс стадийной декомпозиции и "
    "диагностики агентной памяти. Реализует оригинальный метод "
    "четырёхстадийной декомпозиции: Q (запрос) → R (поиск) → M "
    "(материализация цели) → C (завершение). Формальная новизна: "
    "стадийно-марковское условие и аргумент об оценимости стадий из "
    "логов траекторий; мультиагентное расширение с каналами памяти и "
    "занятости. Эмпирическая новизна: сдвиг узкого места M↔C по "
    "каденции K подтверждён на 35 640 прогонах при p<10⁻⁴. "
    "Прикладная новизна: субстрат-независимый Q/R/M/C-логгер + "
    "встроенный прототип Collective Semantic Memory, статистически "
    "строго доминирующий пять peer-broadcast baseline на 3 240 "
    "прогонах. Теоретическое основание: связь Галуа между запросом и "
    "поиском (Formal Concept Analysis)."
)
assert len(QRMC_ABSTRACT_900) <= 900, f"Abstract too long: {len(QRMC_ABSTRACT_900)} chars"

AUTHOR_PLACEHOLDER = "[Ф.И.О. автора — заполнить]"
PERSONAL_PLACEHOLDER = "[заполнить]"
EMAIL_GUESS = "loikoanton@gmail.com"


# ─────────────────────────────────────────────────────────────────────
# 1. Информация для карточки.docx — answers already match Q/R/M/C
# ─────────────────────────────────────────────────────────────────────


def fill_card_info():
    """For Q/R/M/C the applicable answers in the carte are unchanged from
    CADENZA: AI in production (Priority а), trusted software (Critical
    tech 13), AI cross-cutting (Cross-cutting 4). We replace the
    previously chosen answers with refreshed wording to be safe.
    """
    path = f"{ITMO}/Информация для карточки.docx"
    d = Document(path)

    chosen_priority = (
        "а) переход к передовым технологиям проектирования и создания "
        "высокотехнологичной продукции, основанным на применении "
        "интеллектуальных производственных решений, роботизированных и "
        "высокопроизводительных вычислительных систем, новых материалов "
        "и способов конструирования."
    )
    chosen_critical = (
        "13. Технологии создания доверенного и защищенного системного и "
        "прикладного программного обеспечения, в том числе для управления "
        "социальными и экономически значимыми системами."
    )
    chosen_crosscutting = (
        "4. Технологии искусственного интеллекта в отраслях экономики, "
        "социальной сферы (включая сферу общественной безопасности) и в "
        "органах публичной власти."
    )

    # The earlier filler placed answers into the same paragraphs that
    # were originally "Отсутствует / Отсутствуют". The selection is the
    # same here — overwrite to refresh formatting and double-check.
    target_texts = {
        chosen_priority: ("а) переход к передовым технологиям", "priority"),
        chosen_critical: ("13. Технологии создания доверенного", "critical"),
        chosen_crosscutting: ("4. Технологии искусственного интеллекта", "crosscutting"),
    }
    seen = {"priority": False, "critical": False, "crosscutting": False}
    for p in d.paragraphs:
        t = p.text.strip()
        for full, (prefix, key) in target_texts.items():
            if seen[key]:
                continue
            if t.startswith(prefix):
                # Skip the very first occurrence (the enumeration list);
                # the second occurrence is the answer slot.
                seen[key] = "first"
                break
        else:
            continue

    # Run a second pass — replace the second occurrence with the canonical
    # text. This keeps style/font of the existing paragraph.
    counts = {"priority": 0, "critical": 0, "crosscutting": 0}
    for p in d.paragraphs:
        t = p.text.strip()
        for full, (prefix, key) in target_texts.items():
            if t.startswith(prefix):
                counts[key] += 1
                if counts[key] == 2:
                    _replace_paragraph_text(p, full)
                break

    d.save(path)
    print(f"  saved: {path}  (Q/R/M/C uses same priorities/tech as CADENZA, refreshed)")


# ─────────────────────────────────────────────────────────────────────
# 2. Листинг ПР ЭВМ.docx — replace name + replace source code
# ─────────────────────────────────────────────────────────────────────


def fill_listing():
    """Reset the listing to Q/R/M/C: new program name, new source code."""
    path = f"{ITMO}/Листинг ПР ЭВМ.docx"
    d = Document(path)

    # Strategy: walk paragraphs, replace header lines, then drop ALL
    # paragraphs after the "Авторы:" line (those are the old CADENZA
    # source code we want to discard), then append fresh Q/R/M/C code.

    # Step 1: locate the "Авторы:" line — everything below is replaceable.
    cutoff_idx = None
    for i, p in enumerate(d.paragraphs):
        t = p.text.strip()
        if t.startswith("(Название программы)"):
            _replace_paragraph_text(p, QRMC_NAME)
        elif t.startswith("Авторы:"):
            _replace_paragraph_text(p, f"Авторы: {AUTHOR_PLACEHOLDER}")
            cutoff_idx = i

    if cutoff_idx is None:
        raise RuntimeError("Could not locate 'Авторы:' anchor in listing template")

    # Step 2: remove every paragraph after the authors line (old code dump).
    body = d.paragraphs
    for old in list(body[cutoff_idx + 1:]):
        old._element.getparent().remove(old._element)

    # Step 3: append fresh Q/R/M/C source code. Include the runner (the
    # Q/R/M/C logger itself), the analyzer, validators, and supporting
    # modules from experiments/big_experiment/.
    files_to_bundle = [
        ("README.md (experiments/big_experiment/)", f"{BIG_EXP}/README.md"),
        ("experiments/big_experiment/runner.py", f"{BIG_EXP}/runner.py"),
        ("experiments/big_experiment/analyze_qrmc.py", f"{BIG_EXP}/analyze_qrmc.py"),
        ("experiments/big_experiment/env_factory.py", f"{BIG_EXP}/env_factory.py"),
        ("experiments/big_experiment/memory_factory.py", f"{BIG_EXP}/memory_factory.py"),
        ("experiments/big_experiment/planner.py", f"{BIG_EXP}/planner.py"),
        ("experiments/big_experiment/config.py", f"{BIG_EXP}/config.py"),
        ("experiments/big_experiment/exp_cadence_phase.py", f"{BIG_EXP}/exp_cadence_phase.py"),
        ("experiments/big_experiment/exp_oracle_interventions.py", f"{BIG_EXP}/exp_oracle_interventions.py"),
        ("experiments/big_experiment/exp_extra_K.py", f"{BIG_EXP}/exp_extra_K.py"),
        ("experiments/big_experiment/exp_scale_N12.py", f"{BIG_EXP}/exp_scale_N12.py"),
        ("scripts/validate_qrmc_factorization.py", f"{SCRIPTS}/validate_qrmc_factorization.py"),
    ]

    total_lines = 0
    for label, src in files_to_bundle:
        if not os.path.exists(src):
            print(f"  [skip] missing: {src}")
            continue
        d.add_paragraph()
        hdr = d.add_paragraph(f"─── {label} ───────────────────────")
        for r in hdr.runs:
            r.bold = True
        with open(src, encoding="utf-8") as f:
            for line in f.read().splitlines():
                p = d.add_paragraph(line if line else " ")
                _set_mono(p)
                total_lines += 1

    sheets_estimate = max(1, total_lines // 55)
    # Also patch the "Всего листов:" line if it exists earlier in the doc.
    for p in d.paragraphs[:cutoff_idx + 1]:
        if p.text.strip().startswith("Всего листов:"):
            _replace_paragraph_text(p, f"Всего листов: {sheets_estimate}")
            break

    d.save(path)
    print(f"  saved: {path}  (sheets≈{sheets_estimate}, "
          f"{len([f for _, f in files_to_bundle if os.path.exists(f)])} files, "
          f"{total_lines} LOC)")


# ─────────────────────────────────────────────────────────────────────
# 3. Реферат ПР ЭВМ.docx  (rewrite the abstract)
# ─────────────────────────────────────────────────────────────────────


def fill_abstract():
    """Rewrite the abstract for Q/R/M/C. Operates on the .docx we
    already produced in the previous (CADENZA) pass."""
    src_docx = f"{ITMO}/Реферат ПР ЭВМ.docx"
    if not os.path.exists(src_docx):
        # If we lost the .docx version, regenerate from the .doc once.
        src_doc = f"{ITMO}/Реферат ПР ЭВМ.doc"
        subprocess.run(
            ["textutil", "-convert", "docx", "-output", src_docx, src_doc],
            check=True,
        )
    d = Document(src_docx)

    fields = {
        "Программа:": f"Программа: {QRMC_NAME}",
        "Реферат:": f"Реферат: {QRMC_ABSTRACT_900}",
        "Тип ЭВМ:": "Тип ЭВМ: IBM PC-совместимый персональный компьютер",
        "Языки:": "Языки: Python 3.9+",
        "ОС:": "ОС: macOS / Linux / Windows (кросс-платформенно)",
        "Объем программы:": (
            "Объем программы: 95 КБ исходного кода (2 487 строк в 12 модулях)"
        ),
    }

    for p in d.paragraphs:
        for key, full in fields.items():
            if p.text.strip().startswith(key):
                _replace_paragraph_text(p, full)
                break

    d.save(src_docx)
    print(f"  saved: {src_docx}")


# ─────────────────────────────────────────────────────────────────────
# 4. Уведомление о создании РИД.docx — rewrite description + contribution
# ─────────────────────────────────────────────────────────────────────


def fill_notification():
    path = f"{ITMO}/Уведомление_о_создании_РИД_2025.docx"
    d = Document(path)

    # Step 1: rewrite the opening sentence to the new name.
    for p in d.paragraphs:
        if "Настоящим уведомляю" in p.text and "Университет ИТМО" in p.text:
            new = p.text
            # Replace any previous program name between the « and » quotes
            # with the new Q/R/M/C name.
            import re
            new = re.sub(r"«[^»]*»\s*\(название\)",
                         f"«{QRMC_NAME}» (название)", new)
            if "«" not in new:
                # The earlier filler may have left a different format —
                # use a fresh insertion before "(название)".
                new = new.replace(
                    "(название)", f"«{QRMC_NAME}» (название)")
            _replace_paragraph_text(p, new)
            break

    # Step 2: refresh the first big table (Characteristics of RID).
    answers = {
        "Тип РИД (предполагаемый)": "Программа для ЭВМ",
        "Даты начала и окончания создания РИД": (
            "Начало: 01.02.2025. Окончание: 14.06.2026."
        ),
        "Сведения об обнародовании РИД": (
            "Не обнародован. Сопроводительная научная статья находится "
            "в подготовке к подаче на конференцию NeurIPS 2026; на дату "
            "уведомления текст статьи и исходный код в открытый доступ "
            "не передавались."
        ),
        "Сведения об использовании в РИД иных результатов": (
            "При создании РИД использовалось свободное программное "
            "обеспечение, распространяемое под открытыми лицензиями "
            "(Python Software Foundation License — стандартная "
            "библиотека Python 3.9+; BSD-3-Clause — NumPy, SciPy; MIT — "
            "MiniGrid, BabyAI как тестовые окружения). Иные результаты "
            "интеллектуальной деятельности третьих лиц, охраняемые "
            "патентами или авторским правом и закрытыми лицензиями, "
            "при создании РИД не использовались."
        ),
        "Источник финансирования работ по созданию РИД": (
            "Собственные средства автора в рамках диссертационной "
            "работы. Номер проекта ИТМО: [при наличии — заполнить]. "
            "Наименование проекта: диссертационная работа. Заказчик "
            "работ: отсутствует. Номер и дата контракта: отсутствует."
        ),
        "Число экземпляров РИД": (
            "Один экземпляр. Место хранения: репозиторий с исходным "
            "кодом (анонимизированная копия — на USB-носителе и "
            "локальном диске автора по адресу проживания). "
            "Документация о РИД хранится совместно с исходным кодом."
        ),
        "Описание РИД": QRMC_FULL_DESCRIPTION,
    }
    multi_selections = {
        "Возможно ли использование РИД для создания сквозных технологий": [
            "Технология хранения и анализа больших данных",
            "Искусственный интеллект",
        ],
        "Для развития каких рынков Национальной технологической инициативы": [
            "Нейронет",
            "Технет",
        ],
        "Использование результата может обеспечить реализацию приоритетов": [
            "Переход к передовым цифровым, интеллектуальным производственным "
            "технологиям, роботизированным системам, новым материалам и "
            "способам конструирования, создание систем обработки больших "
            "объемов данных, машинного обучения и искусственного интеллекта",
        ],
        "Приоритетное направление развития университета": [
            "Интеллектуальные технологии и робототехника",
            "Информационные технологии в экономике, социальной сфере и искусстве",
        ],
    }

    tbl_props = d.tables[0]
    for row in tbl_props.rows:
        if len(row.cells) < 2:
            continue
        key_cell = row.cells[0].text.strip()
        matched = False
        for prefix, ans in answers.items():
            if key_cell.startswith(prefix):
                _set_cell_text(row.cells[1], ans)
                matched = True
                break
        if matched:
            continue
        for prefix, choices in multi_selections.items():
            if key_cell.startswith(prefix):
                cur = row.cells[1].text
                # If "Выбрано:" already present from the CADENZA pass,
                # preserve the choices (they are the same for Q/R/M/C).
                if cur.strip().startswith("Выбрано:"):
                    break
                new_text = "Выбрано:\n" + "\n".join(f"• {c}" for c in choices)
                new_text += "\n\n— исходный перечень вариантов —\n" + cur
                _set_cell_text(row.cells[1], new_text)
                break

    # Step 3: rewrite the author's creative contribution to be Q/R/M/C-specific.
    qrmc_contribution = (
        "Постановка задачи о стадийной декомпозиции памяти-основанной "
        "навигации; формулировка стадий Q (формирование запроса), "
        "R (удовлетворение поиска), M (материализация цели), "
        "C (завершение после поиска) и стадийно-марковского условия; "
        "формальный аргумент об оценимости стадийных частот из логов "
        "(Definition 1 и Definition 2); формулировка и эмпирическая "
        "верификация утверждения о сдвиге узкого места M↔C по каденции "
        "K (Empirical Claim 1) на 35 640 прогонах при p<10⁻⁴; "
        "построение субстрат-независимого Q/R/M/C-логгера, "
        "работающего на трёх окружениях (MiniGrid, BabyAI, кастомная "
        "сетка); реализация всех 12 модулей программного комплекса; "
        "реализация минимального экземпляра Collective Semantic "
        "Memory и его сравнения с пятью peer-broadcast архитектурами "
        "(3 240 прогонов); теоретическое чтение Q-R-полярности через "
        "связь Галуа в Formal Concept Analysis. Доля вклада: 100%."
    )

    author_updates = {
        "Творческий вклад в создание результата": qrmc_contribution,
    }

    for tbl_idx in (1, 2):
        if tbl_idx >= len(d.tables):
            break
        for row in d.tables[tbl_idx].rows:
            if len(row.cells) < 3:
                continue
            label = row.cells[1].text.strip()
            for prefix, value in author_updates.items():
                if label.startswith(prefix):
                    _set_cell_text(row.cells[2], value)
                    break

    d.save(path)
    print(f"  saved: {path}")


# ─────────────────────────────────────────────────────────────────────
# Helpers — keep paragraph/run styling intact
# ─────────────────────────────────────────────────────────────────────


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    first = paragraph.runs[0]
    first.text = new_text
    for r in list(paragraph.runs[1:]):
        r._element.getparent().remove(r._element)


def _set_cell_text(cell, new_text: str) -> None:
    lines = new_text.split("\n")
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.add_paragraph(lines[0])
        first_p = cell.paragraphs[0]
    else:
        first_p = paragraphs[0]
        _replace_paragraph_text(first_p, lines[0])
        for old in list(paragraphs[1:]):
            old._element.getparent().remove(old._element)
    for line in lines[1:]:
        p = cell.add_paragraph()
        if first_p.style is not None:
            p.style = first_p.style
        if line:
            run = p.add_run(line)
            if first_p.runs and first_p.runs[0].font.name:
                run.font.name = first_p.runs[0].font.name


def _set_mono(paragraph) -> None:
    for r in paragraph.runs:
        r.font.name = "Courier New"
        r.font.size = Pt(9)


def main():
    print("Refilling ITMO RID documents — switching from CADENZA to Q/R/M/C-Диаг …")
    print("=" * 78)
    print("[1/4] Информация для карточки …")
    fill_card_info()
    print("[2/4] Листинг ПР ЭВМ (новый код — experiments/big_experiment/ + validator) …")
    fill_listing()
    print("[3/4] Реферат ПР ЭВМ (новый 900-знаковый текст) …")
    fill_abstract()
    print("[4/4] Уведомление о создании РИД (новые описание + творческий вклад) …")
    fill_notification()
    print()
    print("Done.")
    print()
    print("Личные поля автора по-прежнему `[заполнить]`:")
    print("  ФИО, дата рождения, СНИЛС, ИНН, адрес, телефон,")
    print("  должность, табельный номер, ID РИНЦ/ORCID и т.п.")
    print()
    print(f"Email автора из проектной памяти: {EMAIL_GUESS}")


if __name__ == "__main__":
    main()
