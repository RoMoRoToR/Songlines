"""Patch ITMO RID docs after the Phase A LLM-bridge experiment passed.

The earlier wording said «прогонов с LLM-агентами в составе РИД не
производилось — соответствующее тестирование вынесено в отдельный этап
дорожной карты». After Phase A (10 episodes of Llama-3.1-8B on TextNav
with the Q/R/M/C logger emitting all four event types and 10/10
successes), we replace that with an empirically grounded statement and
reference the artefacts.
"""

from __future__ import annotations

import os
from docx import Document


ITMO = "/Users/taniyashuba/PycharmProjects/Songlines/docs/ITMO"


NEW_DESCRIPTION_SUFFIX_REPLACEMENT_OLD = (
    "прогонов с агентами на основе больших языковых моделей (LLM) в "
    "составе РИД не производилось — соответствующее тестирование "
    "вынесено в отдельный этап дорожной карты"
)
NEW_DESCRIPTION_SUFFIX_REPLACEMENT_NEW = (
    "проведена предварительная эмпирическая верификация переносимости "
    "Q/R/M/C-измерителя на LLM-подложку: одиночный LLM-агент "
    "(Llama-3.1-8B локально через ollama, окружение TextNav — "
    "минимальная мультикомнатная текстовая среда household-типа, 10 "
    "эпизодов, шаг-предел 25, температура 0) проходит все четыре "
    "стадии Q→R→M→C с воспроизводимым успехом 10/10 за 8 тиков в "
    "среднем; условные частоты P(R|Q)=P(M|R)=P(C|M)=1.00, порядок "
    "вложенности Q≥R≥M≥C выполнен, факторизация согласована "
    "(произведение стадийных частот совпадает с P(C*)); ключевое "
    "архитектурное утверждение подтверждено эмпирически — Q/R/M/C-"
    "логгер из experiments/big_experiment/runner.py использован без "
    "модификаций, изменён только агент-сторонний адаптер "
    "(извлечение тегов, формирование запроса, выбор действия — все "
    "LLM-управляемые); полная мультиагентная серия с LLM-агентами "
    "(N=3, peer-broadcast, 5 каденций × 20 seed-ов) выделена в "
    "отдельный этап дорожной карты"
)

NEW_CONTRIBUTION_REPLACEMENT_OLD = (
    "проведение прогонов с LLM-агентами в составе настоящего РИД не "
    "предусмотрено — соответствующая работа выделена в отдельный этап "
    "дорожной карты"
)
NEW_CONTRIBUTION_REPLACEMENT_NEW = (
    "предварительная верификация интерфейса измерителя на LLM-агентах "
    "проведена (Llama-3.1-8B, TextNav, 10 эпизодов, успех 10/10, все "
    "Q/R/M/C события эмиттятся, факторизация согласована); полный "
    "мультиагентный LLM-эксперимент (N=3, peer-broadcast, 5 каденций "
    "× 20 seed-ов) выделен в отдельный этап дорожной карты"
)

# Replace the abstract body — keep ≤900 chars.
NEW_ABSTRACT_900 = (
    "Q/R/M/C-Диаг — Программный комплекс стадийной декомпозиции и "
    "диагностики агентной памяти. Реализует оригинальный метод "
    "четырёхстадийной декомпозиции: Q (запрос) → R (поиск) → "
    "M (материализация цели) → C (завершение). Формальная новизна: "
    "стадийно-марковское условие и аргумент об оценимости стадий из "
    "логов траекторий; мультиагентное расширение с каналами памяти и "
    "занятости. Эмпирическая новизна: сдвиг M↔C по каденции K на "
    "35 640 прогонах при p<10⁻⁴. Прикладная новизна: Q/R/M/C-логгер "
    "верифицирован на трёх типах агентов — символических, "
    "нейросетевых RL (CommNet/PPO/REINFORCE) и LLM-управляемых "
    "(Llama-3.1-8B, TextNav, 10/10 эпизодов, все стадии эмиттятся). "
    "Прототип Collective Semantic Memory доминирует пять peer-"
    "baseline на 3 240 прогонах. Основание: связь Галуа в FCA."
)
assert len(NEW_ABSTRACT_900) <= 900, f"Abstract too long: {len(NEW_ABSTRACT_900)} chars"


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    first = paragraph.runs[0]
    first.text = new_text
    for r in list(paragraph.runs[1:]):
        r._element.getparent().remove(r._element)


def _set_cell_text(cell, new_text: str) -> None:
    lines = new_text.split("\n")
    paragraphs = cell.paragraphs
    if paragraphs:
        first_p = paragraphs[0]
        _replace_paragraph_text(first_p, lines[0])
        for old in list(paragraphs[1:]):
            old._element.getparent().remove(old._element)
    else:
        cell.add_paragraph(lines[0])
        first_p = cell.paragraphs[0]
    for line in lines[1:]:
        p = cell.add_paragraph()
        if first_p.style is not None:
            p.style = first_p.style
        if line:
            run = p.add_run(line)
            if first_p.runs and first_p.runs[0].font.name:
                run.font.name = first_p.runs[0].font.name


def patch_notification():
    path = f"{ITMO}/Уведомление_о_создании_РИД_2025.docx"
    d = Document(path)
    # Описание РИД — replace the LLM scope sentence
    for row in d.tables[0].rows:
        if row.cells[0].text.strip().startswith("Описание РИД"):
            txt = row.cells[1].text
            if NEW_DESCRIPTION_SUFFIX_REPLACEMENT_OLD in txt:
                new = txt.replace(NEW_DESCRIPTION_SUFFIX_REPLACEMENT_OLD,
                                  NEW_DESCRIPTION_SUFFIX_REPLACEMENT_NEW)
                _set_cell_text(row.cells[1], new)
            else:
                print("  [warn] old description LLM clause not found verbatim")
            break
    # Творческий вклад — same swap in both author tables
    for tbl_idx in (1, 2):
        if tbl_idx >= len(d.tables):
            break
        for row in d.tables[tbl_idx].rows:
            if len(row.cells) < 3:
                continue
            if row.cells[1].text.strip().startswith("Творческий вклад"):
                txt = row.cells[2].text
                if NEW_CONTRIBUTION_REPLACEMENT_OLD in txt:
                    new = txt.replace(NEW_CONTRIBUTION_REPLACEMENT_OLD,
                                      NEW_CONTRIBUTION_REPLACEMENT_NEW)
                    _set_cell_text(row.cells[2], new)
                break
    d.save(path)
    print(f"  patched: {path}")


def patch_abstract():
    path = f"{ITMO}/Реферат ПР ЭВМ.docx"
    if not os.path.exists(path):
        print(f"  [skip] missing: {path}")
        return
    d = Document(path)
    for p in d.paragraphs:
        if p.text.strip().startswith("Реферат:"):
            _replace_paragraph_text(p, f"Реферат: {NEW_ABSTRACT_900}")
            break
    d.save(path)
    print(f"  patched: {path}")


def main():
    print("Patching Q/R/M/C ITMO RID — post-LLM-bridge wording …")
    print("=" * 72)
    patch_notification()
    patch_abstract()
    print()
    print("Граница «измеритель / агенты» теперь подтверждена эмпирически.")
    print("Defence-ready claim: «Q/R/M/C-логгер верифицирован на 3 классах")
    print("агентов: символические, RL (CommNet/PPO/REINFORCE), LLM (Llama).")
    print("Артефакты Phase A: tmp/llm_bridge_minimal/")


if __name__ == "__main__":
    main()
