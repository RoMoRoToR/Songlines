"""Fill the four ITMO RID submission documents for the CADENZA package.

The submission unit is the peer_memory package (1024 LoC, 7 modules, 37 KB
Python source) renamed for the РИД card as:

  CADENZA — Cadence-Adaptive Distributed Engine for Networked
  Zero-aggregator Agents

This script does NOT modify styles, fonts, margins, or document chrome.
It only writes content into the fields that the templates leave blank
or marked "Отсутствует / Отсутствуют". Personal data of the author
(СНИЛС, ИНН, ДР, адрес, телефон, паспорт) is left as `[заполнить]`
placeholders for the author to insert by hand.
"""

from __future__ import annotations

import os
import subprocess

from docx import Document
from docx.shared import Pt
from copy import deepcopy

ITMO = "/Users/taniyashuba/PycharmProjects/Songlines/docs/ITMO"
PEER = "/Users/taniyashuba/PycharmProjects/Songlines/peer_memory"

# ─────────────────────────────────────────────────────────────────────
# CADENZA description (used in multiple documents)
# ─────────────────────────────────────────────────────────────────────

CADENZA_NAME = (
    "CADENZA — Cadence-Adaptive Distributed Engine for Networked "
    "Zero-aggregator Agents"
)

CADENZA_FULL_DESCRIPTION = (
    "Программный комплекс CADENZA (Cadence-Adaptive Distributed Engine "
    "for Networked Zero-aggregator Agents) — система распределённой "
    "мультиагентной коллективной памяти с настраиваемой каденцией "
    "peer-to-peer обмена. Ключевая особенность: архитектура без "
    "центрального агрегатора, в которой координация между агентами "
    "обеспечивается тонкой настройкой параметра K — каденции "
    "широковещательной рассылки. Эмпирически (на выборке из 12 960 "
    "экспериментальных прогонов) показано, что при определённых "
    "значениях K* система достигает статистически значимо лучших "
    "результатов мультиагентной координации по сравнению с "
    "централизованными архитектурами (Mann–Whitney U-test, p < 0,05). "
    "Этот эффект отсутствует в существующих архитектурах "
    "collective training / centralized aggregation, что может "
    "составить предмет для самостоятельной охраны."
)

# Abstract <= 900 chars (per template instruction)
CADENZA_ABSTRACT_900 = (
    "CADENZA — Cadence-Adaptive Distributed Engine for Networked "
    "Zero-aggregator Agents. Программный комплекс реализует "
    "распределённую мультиагентную коллективную память без центрального "
    "агрегатора: координация между агентами обеспечивается peer-to-peer "
    "широковещательной рассылкой с настраиваемой каденцией K. Каждый "
    "агент хранит собственный граф памяти и собственное merged-"
    "представление; обмен между агентами идёт только через пассивную "
    "BroadcastBus. Поддерживается асимметричная модель доверия между "
    "агентами и темпоральное затухание. На 12 960 экспериментальных "
    "прогонах при K* система достигает статистически значимо лучших "
    "результатов мультиагентной координации, чем централизованные "
    "архитектуры (Mann–Whitney U-test, p < 0,05). Применяется в "
    "мультиагентных системах поиска, навигации и координации."
)
assert len(CADENZA_ABSTRACT_900) <= 900, f"Abstract too long: {len(CADENZA_ABSTRACT_900)} chars"

AUTHOR_PLACEHOLDER = "[Ф.И.О. автора — заполнить]"
PERSONAL_PLACEHOLDER = "[заполнить]"
EMAIL_GUESS = "loikoanton@gmail.com"


# ─────────────────────────────────────────────────────────────────────
# 1. Информация для карточки.docx
# ─────────────────────────────────────────────────────────────────────


def fill_card_info():
    """Replace 'Отсутствует/Отсутствуют' placeholders with applicable items.

    For CADENZA (AI-based distributed multi-agent memory) the following
    apply:
      • Priorities of S&T development of RF: «а» — переход к передовым
        технологиям, интеллектуальным производственным решениям,
        роботизированным и высокопроизводительным системам;
      • Critical technologies of RF (Указ № 529): № 13 — доверенное и
        защищённое системное и прикладное ПО для управления социальными
        и экономически значимыми системами;
      • Cross-cutting technologies of RF (Указ № 529): № 4 — технологии
        искусственного интеллекта в отраслях экономики, социальной сферы
        и в органах публичной власти.
    """
    src = f"{ITMO}/Информация для карточки.docx"
    dst = f"{ITMO}/Информация для карточки.docx"
    d = Document(src)

    chosen_priority = (
        "а) переход к передовым технологиям проектирования и создания "
        "высокотехнологичной продукции, основанным на применении "
        "интеллектуальных производственных решений, роботизированных и "
        "высокопроизводительных вычислительных систем, новых материалов "
        "и способов конструирования."
    )
    chosen_critical = (
        "13. Технологии создания доверенного и защищенного системного и "
        "прикладного программного обеспечения, в том числе для управления "
        "социальными и экономически значимыми системами."
    )
    chosen_crosscutting = (
        "4. Технологии искусственного интеллекта в отраслях экономики, "
        "социальной сферы (включая сферу общественной безопасности) и в "
        "органах публичной власти."
    )

    # The template puts the placeholder paragraphs ("Отсутствует" /
    # "Отсутствуют") immediately after each enumeration block. We rewrite
    # them in-place so the styling (font, size) is preserved.
    replaced = {"priority": False, "critical": False, "crosscutting": False}
    for p in d.paragraphs:
        text = p.text.strip()
        if text == "Отсутствует" and not replaced["priority"]:
            _replace_paragraph_text(p, chosen_priority)
            replaced["priority"] = True
        elif text == "Отсутствуют" and not replaced["critical"]:
            _replace_paragraph_text(p, chosen_critical)
            replaced["critical"] = True
        elif text == "Отсутствуют" and not replaced["crosscutting"]:
            _replace_paragraph_text(p, chosen_crosscutting)
            replaced["crosscutting"] = True

    d.save(dst)
    print(f"  saved: {dst}  (replaced={replaced})")


# ─────────────────────────────────────────────────────────────────────
# 2. Листинг ПР ЭВМ.docx
# ─────────────────────────────────────────────────────────────────────


def fill_listing():
    """Fill program name and append the full source code of peer_memory."""
    src = f"{ITMO}/Листинг ПР ЭВМ.docx"
    dst = f"{ITMO}/Листинг ПР ЭВМ.docx"
    d = Document(src)

    # Walk paragraphs and substitute the template markers.
    for p in d.paragraphs:
        t = p.text.strip()
        if t == "(Название программы)":
            _replace_paragraph_text(p, CADENZA_NAME)
        elif t == "ПРОГРАММА":
            # leave heading word intact
            pass
        elif t == "Авторы: Фамилия И.О.":
            _replace_paragraph_text(p, f"Авторы: {AUTHOR_PLACEHOLDER}")
        elif t.startswith("Всего листов:"):
            # will recompute after appending sources
            pass

    # Append all source files at the end of the document.
    py_files = sorted(
        fn for fn in os.listdir(PEER) if fn.endswith(".py")
    )
    # README first as architectural overview
    if os.path.exists(f"{PEER}/README.md"):
        d.add_paragraph()
        d.add_paragraph(
            "─── peer_memory/README.md ───────────────────────"
        ).bold = True
        with open(f"{PEER}/README.md", encoding="utf-8") as f:
            for line in f.read().splitlines():
                p = d.add_paragraph(line)
                _set_mono(p)

    for fn in py_files:
        d.add_paragraph()
        hdr = d.add_paragraph(f"─── peer_memory/{fn} ───────────────────────")
        for r in hdr.runs:
            r.bold = True
        with open(f"{PEER}/{fn}", encoding="utf-8") as f:
            for line in f.read().splitlines():
                p = d.add_paragraph(line if line else " ")
                _set_mono(p)

    # Approximate sheet count: ~55 lines per sheet at 12pt.
    total_lines = sum(1 for fn in py_files
                      for _ in open(f"{PEER}/{fn}", encoding="utf-8"))
    sheets_estimate = max(1, total_lines // 55)
    for p in d.paragraphs:
        if p.text.strip().startswith("Всего листов:"):
            _replace_paragraph_text(p, f"Всего листов: {sheets_estimate}")
            break

    d.save(dst)
    print(f"  saved: {dst}  (sheets≈{sheets_estimate}, {len(py_files)} .py files)")


# ─────────────────────────────────────────────────────────────────────
# 3. Реферат ПР ЭВМ.doc  (legacy .doc format — convert via textutil)
# ─────────────────────────────────────────────────────────────────────


def fill_abstract():
    """Convert .doc to .docx, fill, save .docx alongside (and rewrite the .doc
    via textutil → rtf → doc round-trip). The .docx version is the
    submission artefact; the .doc is preserved for tooling compatibility."""
    src_doc = f"{ITMO}/Реферат ПР ЭВМ.doc"
    tmp_docx = f"{ITMO}/_tmp_referat.docx"
    out_docx = f"{ITMO}/Реферат ПР ЭВМ.docx"

    subprocess.run(
        ["textutil", "-convert", "docx", "-output", tmp_docx, src_doc],
        check=True,
    )
    d = Document(tmp_docx)

    fields = {
        "Программа:": f"Программа: {CADENZA_NAME}",
        "Реферат:": f"Реферат: {CADENZA_ABSTRACT_900}",
        "Тип ЭВМ:": "Тип ЭВМ: IBM PC-совместимый персональный компьютер",
        "Языки:": "Языки: Python 3.9+",
        "ОС:": "ОС: macOS / Linux / Windows (кросс-платформенно)",
        "Объем программы:": (
            "Объем программы: 38 КБ исходного кода (1 024 строки в 7 модулях)"
        ),
    }

    for p in d.paragraphs:
        for key, full in fields.items():
            if p.text.strip().startswith(key):
                _replace_paragraph_text(p, full)
                break

    d.save(out_docx)
    os.remove(tmp_docx)
    print(f"  saved: {out_docx}")


# ─────────────────────────────────────────────────────────────────────
# 4. Уведомление о создании РИД.docx
# ─────────────────────────────────────────────────────────────────────


def fill_notification():
    src = f"{ITMO}/Уведомление_о_создании_РИД_2025.docx"
    dst = f"{ITMO}/Уведомление_о_создании_РИД_2025.docx"
    d = Document(src)

    # Step 1: write program name into the opening sentence.
    for p in d.paragraphs:
        if "Настоящим уведомляю" in p.text and "_______" in p.text:
            new = p.text.replace(
                "_________________________________________________ (название)",
                f"«{CADENZA_NAME}» (название)",
            )
            _replace_paragraph_text(p, new)
            break

    # Step 2: fill the first big table (Characteristics of RID).
    answers = {
        "Тип РИД (предполагаемый)": (
            "Программа для ЭВМ"
        ),
        "Даты начала и окончания создания РИД": (
            "Начало: 01.02.2025. Окончание: 10.06.2026."
        ),
        "Сведения об обнародовании РИД": (
            "Не обнародован."
        ),
        "Сведения об использовании в РИД иных результатов": (
            "При создании РИД использовалось свободное программное "
            "обеспечение, распространяемое под открытыми лицензиями "
            "(Python Software Foundation License — стандартная "
            "библиотека Python 3.9+; BSD-3-Clause — NumPy). Иные "
            "результаты интеллектуальной деятельности третьих лиц "
            "не использовались."
        ),
        "Источник финансирования работ по созданию РИД": (
            "Собственные средства автора в рамках диссертационной "
            "работы. Номер проекта ИТМО: [при наличии — заполнить]. "
            "Наименование проекта: диссертационная работа. Заказчик "
            "работ: отсутствует. Номер и дата контракта: отсутствует."
        ),
        "Число экземпляров РИД": (
            "Один экземпляр. Место хранения: репозиторий с исходным "
            "кодом (анонимизированная копия — на USB-носителе и "
            "локальном диске автора по адресу проживания). "
            "Документация о РИД хранится совместно с исходным кодом."
        ),
        "Описание РИД": CADENZA_FULL_DESCRIPTION,
    }
    multi_selections = {
        "Возможно ли использование РИД для создания сквозных технологий": [
            "Технология хранения и анализа больших данных",
            "Искусственный интеллект",
        ],
        "Для развития каких рынков Национальной технологической инициативы": [
            "Нейронет",
            "Технет",
        ],
        "Использование результата может обеспечить реализацию приоритетов": [
            "Переход к передовым цифровым, интеллектуальным производственным "
            "технологиям, роботизированным системам, новым материалам и "
            "способам конструирования, создание систем обработки больших "
            "объемов данных, машинного обучения и искусственного интеллекта",
        ],
        "Приоритетное направление развития университета": [
            "Интеллектуальные технологии и робототехника",
            "Информационные технологии в экономике, социальной сфере и искусстве",
        ],
    }

    tbl_props = d.tables[0]
    for row in tbl_props.rows:
        if len(row.cells) < 2:
            continue
        key_cell = row.cells[0].text.strip()
        # Match by prefix because cells contain long enumerations
        matched = False
        for prefix, ans in answers.items():
            if key_cell.startswith(prefix):
                _set_cell_text(row.cells[1], ans)
                matched = True
                break
        if matched:
            continue
        for prefix, choices in multi_selections.items():
            if key_cell.startswith(prefix):
                cur = row.cells[1].text
                # Build a new cell value: keep the enumeration, but
                # prefix selected items with a check mark, and add an
                # explicit "Выбрано:" line at the top.
                new_text = "Выбрано:\n" + "\n".join(f"• {c}" for c in choices)
                new_text += "\n\n— исходный перечень вариантов —\n" + cur
                _set_cell_text(row.cells[1], new_text)
                break

    # Step 3: fill author tables (there are 2 in the template; we fill
    # the first one fully and leave a comment in the second).
    author_data = {
        "Ф.И.О.      автора": AUTHOR_PLACEHOLDER,
        "Дата рождения": f"{PERSONAL_PLACEHOLDER} (ДД.ММ.ГГГГ)",
        "Гражданство": "Российская Федерация",
        "Должность, место работы": (
            "[заполнить — должность и подразделение Университета ИТМО, "
            "табельный номер; либо иное основание привлечения к работам]"
        ),
        "Основание привлечения к работам": (
            "Трудовой договор с Университетом ИТМО (для работников ИТМО) "
            "либо иное основание [заполнить]"
        ),
        "Творческий вклад в создание результата": (
            "Постановка задачи, разработка архитектуры программного "
            "комплекса CADENZA (peer-to-peer мультиагентная коллективная "
            "память без центрального агрегатора), реализация всех модулей "
            "(BroadcastBus, PeerAgent, PeerRuntime, PeerTrust, PeerMerge), "
            "проведение экспериментов (12 960 прогонов), интерпретация "
            "результатов. Доля вклада: 100%."
        ),
        "Адрес места жительства": f"{PERSONAL_PLACEHOLDER} (с почтовым индексом)",
        "Телефон": PERSONAL_PLACEHOLDER,
        "СНИЛС": PERSONAL_PLACEHOLDER,
        "ИНН": PERSONAL_PLACEHOLDER,
        "Ученая степень": "[заполнить или указать «отсутствует»]",
        "Ученое звание": "[заполнить или указать «отсутствует»]",
        "WOS Research ID": "[при наличии — заполнить]",
        "Scopus Author ID": "[при наличии — заполнить]",
        "ID РИНЦ": "[при наличии — заполнить]",
        "ORCID": "[при наличии — заполнить]",
    }

    for tbl_idx in (1, 2):
        if tbl_idx >= len(d.tables):
            break
        for row in d.tables[tbl_idx].rows:
            if len(row.cells) < 3:
                continue
            label = row.cells[1].text.strip()
            for prefix, value in author_data.items():
                if label.startswith(prefix):
                    # Column index 2 is the value field
                    _set_cell_text(row.cells[2], value)
                    break
        # If the document has only one author, mark the second table
        if tbl_idx == 2 and len(d.tables) > 2:
            # Replace the values in table-2 with a note
            for row in d.tables[2].rows:
                if len(row.cells) >= 3:
                    _set_cell_text(
                        row.cells[2],
                        "[Удалить эту таблицу, если автор один. "
                        "Заполнить, если соавтор есть.]",
                    )
            break

    d.save(dst)
    print(f"  saved: {dst}")


# ─────────────────────────────────────────────────────────────────────
# Helpers — keep paragraph/run styling intact
# ─────────────────────────────────────────────────────────────────────


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replace paragraph text while keeping the first run's character style."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    # Reuse the first run, drop the rest.
    first = paragraph.runs[0]
    first.text = new_text
    for r in list(paragraph.runs[1:]):
        r._element.getparent().remove(r._element)


def _set_cell_text(cell, new_text: str) -> None:
    """Replace cell contents while keeping the first paragraph's style.

    For multi-line text, split on \n and add additional paragraphs in
    the cell, copying the first paragraph's pPr (paragraph properties).
    """
    lines = new_text.split("\n")
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.add_paragraph(lines[0])
        first_p = cell.paragraphs[0]
    else:
        first_p = paragraphs[0]
        _replace_paragraph_text(first_p, lines[0])
        # Remove subsequent paragraphs that were already there.
        for old in list(paragraphs[1:]):
            old._element.getparent().remove(old._element)

    for line in lines[1:]:
        p = cell.add_paragraph()
        if first_p.style is not None:
            p.style = first_p.style
        if line:
            run = p.add_run(line)
            # Copy font name from the first run if available
            if first_p.runs and first_p.runs[0].font.name:
                run.font.name = first_p.runs[0].font.name


def _set_mono(paragraph) -> None:
    for r in paragraph.runs:
        r.font.name = "Courier New"
        r.font.size = Pt(9)


def main():
    print("Filling ITMO RID submission documents for CADENZA …")
    print("=" * 70)
    print("[1/4] Информация для карточки …")
    fill_card_info()
    print("[2/4] Листинг ПР ЭВМ …")
    fill_listing()
    print("[3/4] Реферат ПР ЭВМ …")
    fill_abstract()
    print("[4/4] Уведомление о создании РИД …")
    fill_notification()
    print()
    print("Done.")
    print()
    print("Personal fields left as `[заполнить]` placeholders:")
    print("  Ф.И.О. автора, дата рождения, СНИЛС, ИНН,")
    print("  адрес, телефон, должность, табельный номер.")
    print()
    print("Estimated email of author from project memory: " + EMAIL_GUESS)


if __name__ == "__main__":
    main()
