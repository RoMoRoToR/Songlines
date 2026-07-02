"""Patch the Q/R/M/C ITMO RID documents to make the substrate scope
defense-ready.

The original phrasing «субстрат-независимый Q/R/M/C-логгер, единообразно
работающий на трёх окружениях …» is technically accurate but can be
misread as «работающий с LLM-агентами». This patch tightens the three
slots where the risk lives:

  1. Уведомление о создании РИД — Описание РИД (table 0, row "Описание")
  2. Уведомление о создании РИД — Творческий вклад (tables 1, 2)
  3. Реферат ПР ЭВМ — abstract body (900-char field)

The new wording explicitly distinguishes:
  • Архитектурную независимость интерфейса измерителя от типа агента
    (technical claim about the logger's design);
  • Класс агентов, на которых измеритель эмпирически проверен
    (symbolic planners + RL CommNet/PPO/REINFORCE);
  • Отсутствие прогонов с LLM-агентами и отнесение этой работы к
    отдельному треку дорожной карты.
"""

from __future__ import annotations

import os
from docx import Document

ITMO = "/Users/taniyashuba/PycharmProjects/Songlines/docs/ITMO"


# ─────────────────────────────────────────────────────────────────────
# New text — defense-ready
# ─────────────────────────────────────────────────────────────────────


# Drop-in replacement for the «Описание РИД» cell.
NEW_DESCRIPTION = (
    "Программный комплекс Q/R/M/C-Диаг реализует оригинальный метод "
    "четырёхстадийной декомпозиции работы агентной памяти при "
    "семантической навигации: Q (формирование запроса), R "
    "(удовлетворение поиска), M (материализация цели), C (завершение "
    "после поиска). Метод включает: (а) формальное определение четырёх "
    "стадий через стадийно-марковское условие и аргумент об оценимости "
    "конечных частотных оценок из лога траекторий (Definition 1-2 в "
    "сопроводительной статье); (б) расширение на мультиагентный режим "
    "с двумя каналами связности — памяти и занятости целей "
    "(Definition 2); (в) эмпирически подтверждённое на 35 640 прогонах "
    "утверждение о сдвиге узкого места между стадиями M и C при "
    "изменении каденции K (Empirical Claim 1, оба знака наклонов при "
    "p<10⁻⁴, повторено на 100-прогонном MiniGrid-портабилити-тесте при "
    "p<10⁻⁶); (г) реализацию Q/R/M/C-логгера, входной интерфейс "
    "которого по архитектуре не зависит от внутреннего устройства "
    "агента (логгер читает события из лога траекторий и не использует "
    "информацию о способе принятия решений); эмпирическая верификация "
    "проведена на трёх окружениях семантической навигации (MiniGrid, "
    "BabyAI, кастомная сетка 12×10) с двумя классами агентов: "
    "символическими планировщиками и нейросетевыми RL-baseline "
    "(CommNet с REINFORCE и PPO); прогонов с агентами на основе "
    "больших языковых моделей (LLM) в составе РИД не производилось — "
    "соответствующее тестирование вынесено в отдельный этап дорожной "
    "карты; (д) встроенный прототип Collective Semantic Memory — "
    "минимальный экземпляр распределённой коллективной семантической "
    "памяти с явными правилами merge/trust/staleness, статистически "
    "строго доминирующий пять фиксированных peer-broadcast "
    "архитектур на 3 240 прогонах (нижняя граница 95%-го бутстрап-"
    "доверительного интервала 0,610 превосходит верхнюю границу "
    "каждой из baseline). Теоретическая новизна: операционные "
    "определения Q/R/M/C получают чтение через антитонное "
    "соответствие Галуа между намерением запроса и экстенсией поиска "
    "(концептуальная решётка по Гантеру-Вилле), что превращает "
    "диагноз «retrieval пуст» в строгое утверждение «семантика "
    "запроса лежит вне реализованной концептуальной решётки». "
    "Совокупность формальной структуры, эмпирической верификации на "
    "большой выборке и инструментального применения отсутствует в "
    "существующих архитектурах диагностики мультиагентных систем "
    "памяти и, как предполагается, может составить предмет для "
    "самостоятельной охраны."
)


# Drop-in replacement for the «Творческий вклад» cell in author tables.
NEW_CONTRIBUTION = (
    "Постановка задачи о стадийной декомпозиции памяти-основанной "
    "навигации; формулировка стадий Q (формирование запроса), "
    "R (удовлетворение поиска), M (материализация цели), C (завершение "
    "после поиска) и стадийно-марковского условия; формальный "
    "аргумент об оценимости стадийных частот из логов "
    "(Definition 1 и Definition 2); формулировка и эмпирическая "
    "верификация утверждения о сдвиге узкого места M↔C по каденции K "
    "(Empirical Claim 1) на 35 640 прогонах при p<10⁻⁴; построение "
    "Q/R/M/C-логгера, входной интерфейс которого не зависит от "
    "внутреннего устройства агента (символический планировщик или "
    "нейросетевая политика), с эмпирической верификацией на трёх "
    "окружениях семантической навигации (MiniGrid, BabyAI, кастомная "
    "сетка) и на двух классах агентов (символические + нейросетевые "
    "RL-baseline CommNet/PPO/REINFORCE); проведение прогонов с "
    "LLM-агентами в составе настоящего РИД не предусмотрено — "
    "соответствующая работа выделена в отдельный этап дорожной карты; "
    "реализация всех 12 модулей программного комплекса; реализация "
    "минимального экземпляра Collective Semantic Memory и его "
    "сравнения с пятью peer-broadcast архитектурами (3 240 прогонов); "
    "теоретическое чтение Q-R-полярности через связь Галуа в Formal "
    "Concept Analysis. Доля вклада: 100%."
)


# Drop-in replacement for the Реферат body (must stay ≤900 chars).
NEW_ABSTRACT_900 = (
    "Q/R/M/C-Диаг — Программный комплекс стадийной декомпозиции и "
    "диагностики агентной памяти. Реализует оригинальный метод "
    "четырёхстадийной декомпозиции: Q (запрос) → R (поиск) → "
    "M (материализация цели) → C (завершение). Формальная новизна: "
    "стадийно-марковское условие и аргумент об оценимости стадий из "
    "логов траекторий; мультиагентное расширение с каналами памяти и "
    "занятости. Эмпирическая новизна: сдвиг узкого места M↔C по "
    "каденции K подтверждён на 35 640 прогонах при p<10⁻⁴. "
    "Прикладная новизна: Q/R/M/C-логгер с архитектурно агент-"
    "независимым входом, верифицированный на трёх grid-окружениях с "
    "символическими и RL-агентами (LLM — отдельный трек дорожной "
    "карты); прототип Collective Semantic Memory, доминирующий пять "
    "peer-baseline на 3 240 прогонах. Основание: связь Галуа в FCA."
)
assert len(NEW_ABSTRACT_900) <= 900, f"Abstract too long: {len(NEW_ABSTRACT_900)} chars"


# ─────────────────────────────────────────────────────────────────────
# Apply patches
# ─────────────────────────────────────────────────────────────────────


def patch_notification():
    path = f"{ITMO}/Уведомление_о_создании_РИД_2025.docx"
    d = Document(path)

    # 1) Описание РИД (table 0)
    tbl_props = d.tables[0]
    for row in tbl_props.rows:
        if row.cells[0].text.strip().startswith("Описание РИД"):
            _set_cell_text(row.cells[1], NEW_DESCRIPTION)
            break

    # 2) Творческий вклад (tables 1 and 2)
    for tbl_idx in (1, 2):
        if tbl_idx >= len(d.tables):
            break
        for row in d.tables[tbl_idx].rows:
            if len(row.cells) < 3:
                continue
            if row.cells[1].text.strip().startswith("Творческий вклад"):
                _set_cell_text(row.cells[2], NEW_CONTRIBUTION)
                break

    d.save(path)
    print(f"  patched: {path}")


def patch_abstract():
    path = f"{ITMO}/Реферат ПР ЭВМ.docx"
    if not os.path.exists(path):
        print(f"  [skip] missing: {path}")
        return
    d = Document(path)
    for p in d.paragraphs:
        if p.text.strip().startswith("Реферат:"):
            _replace_paragraph_text(p, f"Реферат: {NEW_ABSTRACT_900}")
            break
    d.save(path)
    print(f"  patched: {path}")


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    first = paragraph.runs[0]
    first.text = new_text
    for r in list(paragraph.runs[1:]):
        r._element.getparent().remove(r._element)


def _set_cell_text(cell, new_text: str) -> None:
    lines = new_text.split("\n")
    paragraphs = cell.paragraphs
    if paragraphs:
        first_p = paragraphs[0]
        _replace_paragraph_text(first_p, lines[0])
        for old in list(paragraphs[1:]):
            old._element.getparent().remove(old._element)
    else:
        cell.add_paragraph(lines[0])
        first_p = cell.paragraphs[0]
    for line in lines[1:]:
        p = cell.add_paragraph()
        if first_p.style is not None:
            p.style = first_p.style
        if line:
            run = p.add_run(line)
            if first_p.runs and first_p.runs[0].font.name:
                run.font.name = first_p.runs[0].font.name


def main():
    print("Patching Q/R/M/C ITMO RID — defense-ready substrate-scope wording …")
    print("=" * 72)
    patch_notification()
    patch_abstract()
    print()
    print("Граница «измеритель / агенты» теперь зафиксирована в трёх местах:")
    print("  • Уведомление → Описание РИД")
    print("  • Уведомление → Творческий вклад автора(-ов)")
    print("  • Реферат ПР ЭВМ → Реферат (900 знаков)")


if __name__ == "__main__":
    main()
